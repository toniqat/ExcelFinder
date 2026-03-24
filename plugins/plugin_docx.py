"""DOCX 파일 플러그인 (python-docx 필요) - Extension Pack"""
import sys
import pandas as pd
from pathlib import Path
from typing import List, Tuple

_src = Path(__file__).parent.parent / "src"
if str(_src) not in sys.path:
    sys.path.insert(0, str(_src))

from plugin_base import FormatPlugin


class DocxPlugin(FormatPlugin):
    """Word 문서(.docx) 처리 플러그인 (python-docx 필요)"""

    @property
    def plugin_id(self) -> str:
        return "docx"

    @property
    def display_name(self) -> str:
        return "Word (DOCX)"

    @property
    def required_packages(self) -> List[str]:
        return ["docx"]

    def supported_extensions(self) -> Tuple[str, ...]:
        return ('.docx', '.doc')

    def read_file(self, file_path: str) -> List[Tuple[str, List[str], pd.DataFrame]]:
        import docx  # lazy import

        ext = Path(file_path).suffix.lower()
        if ext == '.doc':
            try:
                doc = docx.Document(file_path)
            except Exception:
                raise ValueError(
                    f".doc (구형 바이너리 Word) 형식은 python-docx로 읽을 수 없습니다: {file_path}\n"
                    "파일을 .docx 형식으로 변환한 후 다시 시도하세요."
                )
        else:
            doc = docx.Document(file_path)
        file_stem = Path(file_path).stem
        sections = []

        # 테이블 처리
        for t_idx, table in enumerate(doc.tables, 1):
            rows_data = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])

            if not rows_data:
                continue

            # 첫 행을 헤더로
            headers = rows_data[0]
            data_rows = rows_data[1:] if len(rows_data) > 1 else []

            if not data_rows:
                # 헤더만 있는 경우 헤더도 데이터로 포함
                data_rows = [rows_data[0]]
                headers = [f"Col_{i}" for i in range(len(rows_data[0]))]

            # 컬럼 수 맞추기
            n_cols = len(headers)
            normalized = []
            for row in data_rows:
                row_padded = list(row) + [""] * max(0, n_cols - len(row))
                normalized.append(row_padded[:n_cols])

            df = pd.DataFrame(normalized, columns=range(n_cols))
            sections.append((f"Table_{t_idx}", headers, df))

        # 단락(텍스트) 처리
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            headers = ['Content']
            data = [(text,) for text in paragraphs]
            df = pd.DataFrame(data, columns=range(1))
            sections.append((file_stem, headers, df))

        return sections
